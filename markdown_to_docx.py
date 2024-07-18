#!/usr/bin/env python3
import errno
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import markdown2
from bs4 import BeautifulSoup
from PIL import Image
from docx.oxml.shared import OxmlElement, qn
from html.parser import HTMLParser
import re
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def write_out_html(file_name, text_html, encoding='utf8'):
    with open(file_name, 'w', encoding=encoding) as output_fd:
        output_fd.write(text_html)

def find_page_width(doc):
    return float(doc.sections[0].page_width / 914400)

def do_table_of_contents(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:t')
    fld_char3.text = "Right-click to update field."
    fld_char2.append(fld_char3)
    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)

def do_table(doc, table_in, style):
    the_header = table_in.find('thead')
    the_column_names = the_header.find_all('th') if the_header else []
    the_data = table_in.find_all('td')
    n_cols = len(the_column_names) if the_column_names else len(table_in.find('tr').find_all(['td', 'th']))
    n_rows = len(table_in.find_all('tr'))
    this_table = doc.add_table(rows=n_rows, cols=n_cols, style=style)
    
    for i, row in enumerate(table_in.find_all('tr')):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            this_table.cell(i, j).text = cell.get_text(strip=True)
    
    # Make the header row bold
    for cell in this_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

def find_image_size(image_file):
    return Image.open(image_file).size

def do_paragraph(line, doc, page_width_inches, style_body, assumed_pixels_per_inch=200, picture_fraction_of_width=0.7):
    is_image = line.find('img')
    if is_image is not None:
        image_source = is_image['src']
        w, h = find_image_size(image_source)
        w_in_inches = w / assumed_pixels_per_inch
        picture_width_inches = page_width_inches * picture_fraction_of_width
        chosen_width = min(picture_width_inches, w_in_inches)
        doc.add_picture(image_source, width=docx.shared.Inches(chosen_width))
        return
    
    paragraph = doc.add_paragraph(style=style_body)
    for child in line.children:
        if child.name == 'strong':
            paragraph.add_run(child.text).bold = True
        elif child.name == 'em':
            paragraph.add_run(child.text).italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif child.name == 'a':
            add_hyperlink(paragraph, child['href'], child.text)
        else:
            paragraph.add_run(child.text)

def do_pre_code(line, doc, style_quote_table):
    table = doc.add_table(rows=1, cols=1, style=style_quote_table)
    cell = table.cell(0, 0)
    cell.text = line.text.strip()
    paragraphs = cell.paragraphs
    paragraph = paragraphs[0]
    run_obj = paragraph.runs
    run = run_obj[0]
    font = run.font
    font.size = Pt(10)
    font.name = 'Courier New'

def do_horizontal_rule(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run('─' * 50)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(192, 192, 192)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def do_code_block(doc, code_text, style_code_block):
    paragraph = doc.add_paragraph()
    paragraph.style = style_code_block
    run = paragraph.add_run(code_text)
    font = run.font
    font.name = 'Courier New'
    font.size = Pt(9)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.right_indent = Inches(0.5)
    shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
    paragraph._p.get_or_add_pPr().append(shading_elm)

class HtmlListParser(HTMLParser):
    list_level = -1
    lists = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
    ordered_lists = ['List Number', 'List Number 2', 'List Number 3']
    doc = None
    spacing = '    '
    spare_list = '○  '
    current_list_type = None

    def handle_starttag(self, tag, attrs):
        if tag in ['ol', 'ul']:
            self.list_level += 1
            self.current_list_type = 'ol' if tag == 'ol' else 'ul'

    def handle_endtag(self, tag):
        if tag in ['ol', 'ul']:
            self.list_level -= 1

    def handle_data(self, data):
        data = data.strip()
        if data:
            if self.list_level in range(len(self.lists)):
                style = self.ordered_lists[self.list_level] if self.current_list_type == 'ol' else self.lists[self.list_level]
                self.doc.add_paragraph(data, style=style)
            else:
                self.doc.add_paragraph('        ' + self.spacing * self.list_level + self.spare_list + data)

class Markdown2docx:
    style_table = 'Medium Shading 1 Accent 3'
    style_quote = 'Intense Quote'
    style_body = 'Body Text'
    style_quote_table = 'Table Grid'
    style_code_block = 'Code Block'
    toc_indicator = 'contents'

    def __init__(self, project, markdown_content):
        self.project = project
        self.outfile = f"{project}.docx"
        self.html_out_file = f"{project}.html"
        self.doc = docx.Document()
        self.page_width_inches = find_page_width(self.doc)
        self.markdown = markdown_content
        self.html = markdown2.markdown(self.markdown, extras=[
            'fenced-code-blocks',
            'code-friendly',
            'wiki-tables',
            'tables',
            'break-on-newline'
        ])
        self.soup = BeautifulSoup(self.html, 'html.parser')
        self.create_code_block_style()

    def create_code_block_style(self):
        styles = self.doc.styles
        style = styles.add_style(self.style_code_block, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(9)
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.right_indent = Inches(0.5)

    def eat_soup(self):
        table_of_contents_done = 0
        for line in self.soup.children:
            if isinstance(line, str):
                continue
            if str(line).lower().find(self.toc_indicator) >= 0 and table_of_contents_done < 2:
                table_of_contents_done += 1
                if table_of_contents_done == 2:
                    do_table_of_contents(self.doc)
            elif line.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(line.name[1]) - 1
                self.doc.add_heading(line.text.strip(), level)
            elif line.name == 'p':
                do_paragraph(line, self.doc, self.page_width_inches, self.style_body)
            elif line.name == 'pre':
                code_block = line.find('code')
                if code_block:
                    do_code_block(self.doc, code_block.get_text(), self.style_code_block)
                else:
                    do_pre_code(line, self.doc, self.style_quote_table)
            elif line.name == 'blockquote':
                self.doc.add_paragraph(line.text.strip(), style=self.style_quote)
            elif line.name == 'hr':
                do_horizontal_rule(self.doc)
            elif line.name == 'table':
                do_table(self.doc, line, self.style_table)
            elif line.name in ['ul', 'ol']:
                parser = HtmlListParser()
                parser.doc = self.doc
                parser.feed(str(line))

    def write_html(self):
        write_out_html(self.html_out_file, self.html)

    def save(self):
        self.doc.save(self.outfile)

def markdown_to_docx(project_name, markdown_content):
    project = Markdown2docx(project_name, markdown_content)
    project.eat_soup()
    project.write_html()  # optional
    project.save()
    return project.outfile

if __name__ == "__main__":
    # Example usage
    project_name = "example_project"
    markdown_content = """
    # Hello, Markdown!

    This is a sample markdown content.

    ## Features
    - Easy to use
    - Converts to Word document

    1. Numbered list
    2. With multiple items

    > This is a blockquote

    ```python
    def hello_world():
        print("Hello, World!")
    ```

    [Link to Google](https://www.google.com)

    ---

    | Column 1 | Column 2 |
    |----------|----------|
    | Cell 1   | Cell 2   |
    """
    output_file = markdown_to_docx(project_name, markdown_content)
    print(f"Docx file created: {output_file}")
